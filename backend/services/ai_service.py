# backend/services/ai_service.py - 文檔管理 AI 服務（v2 — 全面優化版）
# -*- coding: utf-8 -*-
#
# 優化清單:
#   P0  修正 async/await bug（generate 函式正確 await）
#   P1  Embedding 模型可配置，預設 multilingual（中英雙語）
#   P2  RRF (Reciprocal Rank Fusion) 取代 len() 排序
#   P3  結構感知切塊 + Excel 全量擷取
#   P4  Prompt 強化（anti-hallucination + 結構化輸出）
#   P5  HyDE 條件觸發（首次 score 低於閾值才啟用）
#   P6  FAISS IndexFlatIP + L2 normalize
#   P7  Parent-Child pickle 快取（加速重啟）
#   P8  DocumentManager 連線改進 → PostgreSQL

import asyncio
import os
import pickle
import requests
import psycopg2
import psycopg2.extras
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
import pandas as pd
from pathlib import Path
import hashlib
import logging
import re
import shutil
import threading
import warnings

# 精準抑制特定警告
warnings.filterwarnings(
    "ignore",
    message="`encoder_attention_mask` is deprecated and will be removed in version 4.55.0",
    category=FutureWarning,
)
warnings.filterwarnings(
    "ignore",
    message="Valid config keys have changed in V2",
    category=UserWarning,
)

# ─────────────────────────── 檔案/OCR 依賴 ───────────────────────────
import PyPDF2
import docx
import openpyxl  # noqa: F401
from PIL import Image
import pytesseract

# ─────────────────────────── LangChain / 向量庫 ───────────────────────────
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.schema import Document

from langchain.retrievers import ParentDocumentRetriever, ContextualCompressionRetriever
from langchain.storage import InMemoryStore
from langchain.retrievers.document_compressors import (
    EmbeddingsFilter,
    DocumentCompressorPipeline,
)

try:
    from langchain.retrievers.contextual_compression import LLMChainExtractor
    from langchain_openai import ChatOpenAI
except Exception:
    LLMChainExtractor = None
    ChatOpenAI = None

import faiss
import numpy as np
from langchain_community.docstore import InMemoryDocstore

# ─────────────────────────── PostgreSQL 連線 ───────────────────────────
from core.pg import get_conn, get_cursor

SCHEMA = "documents"

# ─────────────────────────── 基本設定 ───────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("services.ai_service")

DOCUMENTS_DIR = Path("documents")
DOCUMENTS_DIR.mkdir(exist_ok=True)

INDEX_DIR = Path("vectorstore")
INDEX_DIR.mkdir(exist_ok=True)

PARENT_CHILD_CACHE = INDEX_DIR / "parent_child.pkl"

USE_FTS = True

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md",
    ".xlsx", ".xls", ".csv",
    ".png", ".jpg", ".jpeg", ".tiff"
}

DOCUMENT_CATEGORIES = {
    "sop": "Standard Operating Procedures",
    "form": "Forms and Templates",
    "manual": "Manuals and Guides",
    "policy": "Policies and Regulations",
    "checklist": "Checklists",
    "other": "Other Documents",
}

DEFAULT_OLLAMA_MODEL = "gpt-oss:20b"

# ── P1: 可配置 Embedding 模型 ──
# 推薦升級路徑（準確率由低到高）:
#   sentence-transformers/all-MiniLM-L6-v2          (384d, 80MB,  英文)
#   sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2  (384d, 471MB, 多語言)
#   BAAI/bge-base-zh-v1.5                           (768d, 400MB, 中英)
#   BAAI/bge-m3                                     (1024d, 2.2GB, 最強多語言)
#
# 變更模型後須執行 /api/ai/reindex 重建向量庫
EMBEDDING_MODEL = os.getenv(
    "EMBEDDING_MODEL",
    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
)

# ── P2: RRF 常數 ──
RRF_K = 60  # Reciprocal Rank Fusion 常數

# ── P5: HyDE 閾值 — 首次檢索最高分低於此值才觸發 HyDE ──
HYDE_SCORE_THRESHOLD = float(os.getenv("HYDE_SCORE_THRESHOLD", "0.45"))
# 文件數低於此值時關閉 HyDE（ROI 太低）
HYDE_MIN_DOCS = int(os.getenv("HYDE_MIN_DOCS", "5"))


# ═══════════════════════════════════════════════════════════════════
# DocumentManager（文件/DB/擷取/切塊/FTS）
# ═══════════════════════════════════════════════════════════════════
class DocumentManager:
    """文檔管理：PostgreSQL + 檔案 CRUD + 文字擷取 + 切塊 + FTS"""

    def __init__(self):
        # PostgreSQL always has full-text search via tsvector
        self.fts_enabled = True
        logger.info("DocumentManager initialized (PostgreSQL + tsvector FTS)")

    # ────────────── 基礎工具 ──────────────
    def get_file_hash(self, file_path: str) -> str:
        h = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()

    # ────────────── 文字擷取 ──────────────
    def extract_text_from_file(self, file_path: str, file_type: str) -> tuple[str, str]:
        try:
            if file_type == ".pdf":
                return self._extract_from_pdf(file_path)
            elif file_type in [".docx", ".doc"]:
                return self._extract_from_docx(file_path)
            elif file_type in [".txt", ".md"]:
                return self._extract_from_text(file_path)
            elif file_type in [".xlsx", ".xls", ".csv"]:
                return self._extract_from_excel(file_path)
            elif file_type in [".png", ".jpg", ".jpeg", ".tiff"]:
                return self._extract_from_image(file_path)
            else:
                return "", f"Unsupported file type: {file_type}"
        except Exception as e:
            logger.exception(f"Error extracting text: {file_path}")
            return "", f"Error: {str(e)}"

    def _extract_from_pdf(self, file_path: str) -> tuple[str, str]:
        try:
            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    ptxt = page.extract_text() or ""
                    if not ptxt.strip() and i < 5:
                        try:
                            from pdf2image import convert_from_path
                            imgs = convert_from_path(file_path, first_page=i + 1, last_page=i + 1, dpi=200)
                            if imgs:
                                ptxt = pytesseract.image_to_string(imgs[0], lang="eng+chi_tra+chi_sim")
                        except Exception:
                            pass
                    text += ptxt + "\n"
            return text.strip(), ""
        except Exception as e:
            return "", f"PDF extraction error: {str(e)}"

    def _extract_from_docx(self, file_path: str) -> tuple[str, str]:
        try:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return text.strip(), ""
        except Exception as e:
            return "", f"DOCX extraction error: {str(e)}"

    def _extract_from_text(self, file_path: str) -> tuple[str, str]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip(), ""
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                    return f.read().strip(), ""
            except Exception as e:
                return "", f"Text extraction error: {str(e)}"
        except Exception as e:
            return "", f"Text extraction error: {str(e)}"

    def _extract_from_excel(self, file_path: str) -> tuple[str, str]:
        """P3: 全量擷取 Excel/CSV — 每個 sheet 獨立、包含欄位 metadata"""
        try:
            if file_path.endswith(".csv"):
                try:
                    df = pd.read_csv(file_path)
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, encoding="utf-8", errors="ignore")
                sheets = {"Sheet1": df}
            else:
                xls = pd.ExcelFile(file_path)
                sheets = {name: xls.parse(name) for name in xls.sheet_names}

            parts: list[str] = [f"File: {os.path.basename(file_path)}"]
            for sheet_name, df in sheets.items():
                if df.empty:
                    continue
                parts.append(f"\n--- Sheet: {sheet_name} ---")
                parts.append(f"Columns: {', '.join(map(str, df.columns.tolist()))}")
                parts.append(f"Rows: {len(df)}")
                # 全量轉文字（超過 500 行時截斷並標註）
                if len(df) > 500:
                    parts.append(df.head(500).to_string(index=False))
                    parts.append(f"... ({len(df) - 500} more rows truncated)")
                else:
                    parts.append(df.to_string(index=False))
            return "\n".join(parts), ""
        except Exception as e:
            return "", f"Excel extraction error: {str(e)}"

    def _extract_from_image(self, file_path: str) -> tuple[str, str]:
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang="eng+chi_tra+chi_sim")
            return text.strip(), ""
        except Exception as e:
            return "", f"OCR extraction error: {str(e)}"

    # ────────────── 切塊 ──────────────
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """P3: 結構感知切塊 — 先按 Markdown 標題/段落分，再按大小切"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=64,
            separators=[
                "\n## ", "\n### ", "\n# ",     # Markdown headings
                "\n\n",                         # Paragraphs
                "\n",                           # Lines
                ". ", "。", "；",               # Sentences
                " ", "",                        # Words/chars
            ],
        )
        return splitter.split_text(text)

    # ────────────── 上傳 / 刪除 / 讀取 ──────────────
    def upload_document(
        self,
        file_path: str,
        original_name: str,
        category: str,
        uploaded_by: str,
        tags: str = "",
        description: str = "",
    ) -> Dict[str, Any]:
        try:
            if not os.path.exists(file_path):
                return {"success": False, "error": "File not found"}

            file_size = os.path.getsize(file_path)
            file_type = Path(file_path).suffix.lower()
            if file_type not in SUPPORTED_EXTENSIONS:
                return {"success": False, "error": f"Unsupported file type: {file_type}"}

            file_hash = self.get_file_hash(file_path)

            with get_cursor(SCHEMA) as cur:
                cur.execute("SELECT id FROM documents WHERE file_hash=%s", (file_hash,))
                exists = cur.fetchone()
            if exists:
                return {"success": False, "error": "Document already exists"}

            content, error = self.extract_text_from_file(file_path, file_type)
            if error and not content:
                return {"success": False, "error": error}

            preview = content[:500] + "..." if len(content) > 500 else content
            if tags:
                tags = ",".join(sorted(set(t.strip().lower() for t in tags.split(",") if t.strip())))

            chunks = self._split_text_into_chunks(content) if content else []
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = re.sub(r"[^\w\-_\.]", "_", original_name)
            new_filename = f"{timestamp}_{safe_name}"
            new_file_path = DOCUMENTS_DIR / new_filename

            with get_conn(SCHEMA) as conn:
                cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                cur.execute("""
                    INSERT INTO documents
                    (filename, original_name, category, file_type, file_size, file_hash,
                     content_preview, uploaded_by, tags, description)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (new_filename, original_name, category, file_type, file_size,
                      file_hash, preview, uploaded_by, tags, description))
                document_id = cur.fetchone()["id"]

                for i, ck in enumerate(chunks):
                    cur.execute("""
                        INSERT INTO document_chunks (document_id, chunk_index, content, chunk_size)
                        VALUES (%s, %s, %s, %s)
                    """, (document_id, i, ck, len(ck)))

            shutil.copy2(file_path, new_file_path)
            return {
                "success": True,
                "document_id": document_id,
                "filename": new_filename,
                "content_length": len(content),
                "chunks_created": len(chunks),
            }
        except Exception as e:
            logger.exception("Error uploading document")
            return {"success": False, "error": str(e)}

    def get_documents(self, category: str = None, status: str = "active") -> List[Dict[str, Any]]:
        with get_cursor(SCHEMA) as cur:
            if category:
                cur.execute(
                    "SELECT * FROM documents WHERE status=%s AND category=%s ORDER BY upload_date DESC",
                    (status, category),
                )
            else:
                cur.execute(
                    "SELECT * FROM documents WHERE status=%s ORDER BY upload_date DESC",
                    (status,),
                )
            rows = cur.fetchall()
        return [{
            "id": r["id"],
            "filename": r["filename"],
            "original_name": r["original_name"],
            "category": r["category"],
            "file_type": r["file_type"],
            "file_size": r["file_size"],
            "content_preview": r["content_preview"],
            "upload_date": r["upload_date"],
            "uploaded_by": r["uploaded_by"],
            "tags": r["tags"],
            "description": r["description"],
        } for r in rows]

    def get_document_by_id(self, document_id: int) -> Optional[Dict[str, Any]]:
        with get_cursor(SCHEMA) as cur:
            cur.execute(
                "SELECT * FROM documents WHERE id=%s AND status='active'", (document_id,)
            )
            r = cur.fetchone()
        if not r:
            return None
        return {
            "id": r["id"],
            "filename": r["filename"],
            "original_name": r["original_name"],
            "category": r["category"],
            "file_type": r["file_type"],
            "file_size": r["file_size"],
            "file_hash": r["file_hash"],
            "content_preview": r["content_preview"],
            "upload_date": r["upload_date"],
            "uploaded_by": r["uploaded_by"],
            "tags": r["tags"],
            "description": r["description"],
        }

    def search_documents_by_name(self, search_term: str) -> List[Dict[str, Any]]:
        rows = []
        if self.fts_enabled and search_term.strip():
            try:
                q = self._fts_safe_query(search_term)
                with get_cursor(SCHEMA) as cur:
                    cur.execute("""
                        SELECT d.*
                        FROM documents d
                        WHERE d.status='active'
                          AND (to_tsvector('simple', d.original_name || ' ' || COALESCE(d.tags, '') || ' ' || COALESCE(d.description, ''))
                               @@ plainto_tsquery('simple', %s))
                        ORDER BY ts_rank(
                            to_tsvector('simple', d.original_name || ' ' || COALESCE(d.tags, '') || ' ' || COALESCE(d.description, '')),
                            plainto_tsquery('simple', %s)
                        ) DESC
                        LIMIT 50
                    """, (q, q))
                    rows = cur.fetchall()
            except Exception as e:
                logger.warning(f"FTS 文件查詢失敗，使用 LIKE：{e}")
                rows = []
        if not rows:
            like = f"%{search_term.lower()}%"
            with get_cursor(SCHEMA) as cur:
                cur.execute("""
                    SELECT * FROM documents
                    WHERE status='active' AND (
                        LOWER(original_name) LIKE %s OR LOWER(tags) LIKE %s OR LOWER(description) LIKE %s
                    )
                    ORDER BY upload_date DESC
                    LIMIT 50
                """, (like, like, like))
                rows = cur.fetchall()

        return [{
            "id": r["id"],
            "filename": r["filename"],
            "original_name": r["original_name"],
            "category": r["category"],
            "file_type": r["file_type"],
            "file_size": r["file_size"],
            "upload_date": r["upload_date"],
            "tags": r["tags"],
            "description": r["description"],
        } for r in rows]

    def get_document_path(self, document_id: int) -> Optional[Path]:
        doc = self.get_document_by_id(document_id)
        if doc:
            return DOCUMENTS_DIR / doc["filename"]
        return None

    def delete_document(self, document_id: int) -> bool:
        try:
            with get_cursor(SCHEMA) as cur:
                cur.execute("SELECT filename FROM documents WHERE id=%s", (document_id,))
                r = cur.fetchone()
            if not r:
                return False
            filename = r["filename"]

            with get_conn(SCHEMA) as conn:
                cur = conn.cursor()
                cur.execute("DELETE FROM document_chunks WHERE document_id=%s", (document_id,))
                cur.execute("DELETE FROM documents WHERE id=%s", (document_id,))

            f = DOCUMENTS_DIR / filename
            if f.exists():
                os.remove(f)
            return True
        except Exception as e:
            logger.exception(f"Error deleting document {document_id}")
            return False

    def get_full_text_by_document_id(self, document_id: int) -> str:
        with get_cursor(SCHEMA) as cur:
            cur.execute("""
                SELECT content FROM document_chunks
                WHERE document_id=%s
                ORDER BY chunk_index ASC
            """, (document_id,))
            rows = cur.fetchall()
        return "\n".join([r["content"] for r in rows]) if rows else ""

    def _fts_safe_query(self, text: str) -> str:
        """Sanitize query for PostgreSQL plainto_tsquery — just return cleaned terms."""
        toks = re.findall(r"[A-Za-z0-9\u4e00-\u9fff]+", text.lower())
        return " ".join(toks) if toks else text.strip()


# ═══════════════════════════════════════════════════════════════════
# RAG 核心：FAISS + RRF + Parent-Child + HyDE（條件式）+ 壓縮
# ═══════════════════════════════════════════════════════════════════
class DocumentRAG:

    def __init__(self, document_manager: DocumentManager):
        self.dm = document_manager
        self.embedding_model = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            encode_kwargs={"normalize_embeddings": True},  # P6: L2 normalize
        )

        self.vectorstore: Optional[FAISS] = None

        # Parent-Child 結構
        self.parent_store = InMemoryStore()
        self.pc_vectorstore: Optional[FAISS] = None
        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000, chunk_overlap=150,
            separators=["\n## ", "\n### ", "\n# ", "\n\n", "\n", ". ", "。", " ", ""],
        )
        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400, chunk_overlap=60,
            separators=["\n## ", "\n### ", "\n# ", "\n\n", "\n", ". ", "。", " ", ""],
        )
        self.parent_retriever: Optional[ParentDocumentRetriever] = None

        self._load_or_build()

    # ───── P6: 建立 FAISS（IndexFlatIP + normalize） ─────
    def _new_empty_faiss(self) -> FAISS:
        dim = len(self.embedding_model.embed_query("dimension probe"))
        index = faiss.IndexFlatIP(dim)  # Inner Product（因 embeddings 已 L2 normalize）
        return FAISS(
            embedding_function=self.embedding_model,
            index=index,
            docstore=InMemoryDocstore(),
            index_to_docstore_id={},
        )

    # ───── 主 FAISS ─────
    def _load_or_build(self):
        try:
            faiss_path = INDEX_DIR / "index.faiss"
            if faiss_path.exists():
                self.vectorstore = FAISS.load_local(
                    str(INDEX_DIR), self.embedding_model, allow_dangerous_deserialization=True
                )
                # 檢查維度是否匹配（模型變更後需重建）
                expected_dim = len(self.embedding_model.embed_query("dim check"))
                actual_dim = self.vectorstore.index.d
                if expected_dim != actual_dim:
                    logger.warning(
                        f"Warning FAISS 維度不匹配（index={actual_dim}, model={expected_dim}），重建索引…"
                    )
                    self._build_vector_store()
                    self._save()
                else:
                    logger.info("已從磁碟載入主 FAISS")
            else:
                self._build_vector_store()
                self._save()
        except Exception as e:
            logger.error(f"FAISS 載入失敗，改為重建：{e}")
            self._build_vector_store()
            self._save()

        # P7: Parent-Child — 嘗試從 pickle 載入，否則重建
        try:
            self._load_or_build_parent_child()
            logger.info("Parent-Child 結構已建立")
        except Exception as e:
            logger.warning(f"Warning Parent-Child 建立失敗：{e}")

    def _save(self):
        if self.vectorstore:
            self.vectorstore.save_local(str(INDEX_DIR))

    def _build_vector_store(self, exclude_doc_ids: Optional[set[int]] = None):
        exclude_doc_ids = exclude_doc_ids or set()
        with get_cursor(SCHEMA) as cur:
            cur.execute("""
                SELECT dc.id AS chunk_id, dc.chunk_index, dc.content,
                       d.original_name, d.category, d.id AS doc_id, d.tags, d.description
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
                WHERE d.status='active'
            """)
            rows = cur.fetchall()

        docs: List[Document] = []
        for r in rows:
            if r["doc_id"] in exclude_doc_ids:
                continue
            docs.append(Document(
                page_content=r["content"],
                metadata={
                    "document_id": r["doc_id"],
                    "chunk_id": r["chunk_id"],
                    "chunk_index": r["chunk_index"],
                    "filename": r["original_name"],
                    "category": r["category"],
                    "tags": r["tags"] or "",
                    "description": r["description"] or "",
                },
            ))
        if docs:
            self.vectorstore = FAISS.from_documents(docs, self.embedding_model)
            logger.info(f"主 FAISS 重建完成，chunks={len(docs)}")
        else:
            self.vectorstore = None
            logger.warning("Warning 無可用文件建立主向量庫")

    def rebuild_vector_store(self):
        self._build_vector_store()
        self._save()
        try:
            self._build_parent_child_index()
            self._save_parent_child()
        except Exception as e:
            logger.warning(f"Parent-Child 重建失敗：{e}")

    def add_document_chunks(self, document_id: int):
        with get_cursor(SCHEMA) as cur:
            cur.execute("""
                SELECT dc.id AS chunk_id, dc.chunk_index, dc.content,
                       d.original_name, d.category, d.id AS doc_id, d.tags, d.description
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
                WHERE d.status='active' AND d.id=%s
            """, (document_id,))
            rows = cur.fetchall()
        docs = [Document(
            page_content=r["content"],
            metadata={
                "document_id": r["doc_id"],
                "chunk_id": r["chunk_id"],
                "chunk_index": r["chunk_index"],
                "filename": r["original_name"],
                "category": r["category"],
                "tags": r["tags"] or "",
                "description": r["description"] or "",
            },
        ) for r in rows]
        if docs:
            if not self.vectorstore:
                self.vectorstore = FAISS.from_documents(docs, self.embedding_model)
            else:
                self.vectorstore.add_documents(docs)
            self._save()

        try:
            self._add_parent_child_for_doc(document_id)
            self._save_parent_child()
        except Exception as e:
            logger.warning(f"Parent-Child 增量加入失敗（doc={document_id}）：{e}")

    def remove_document(self, document_id: int):
        logger.info(f"重新建立主向量庫（排除文件 {document_id}）")
        self._build_vector_store(exclude_doc_ids={document_id})
        self._save()
        try:
            self._build_parent_child_index()
            self._save_parent_child()
        except Exception as e:
            logger.warning(f"Parent-Child 重建失敗：{e}")

    # ───── P7: Parent-Child pickle 快取 ─────
    def _save_parent_child(self):
        """將 parent store 序列化到磁碟，加速下次啟動"""
        if self.pc_vectorstore is None:
            return
        try:
            data = {
                "store_dict": dict(self.parent_store.yield_keys()),
            }
            # InMemoryStore 內部用 dict；直接存 mget 全部
            all_keys = list(self.parent_store.yield_keys())
            all_vals = self.parent_store.mget(all_keys)
            data = {"keys": all_keys, "values": all_vals}
            with open(PARENT_CHILD_CACHE, "wb") as f:
                pickle.dump(data, f)
            if self.pc_vectorstore:
                self.pc_vectorstore.save_local(str(INDEX_DIR / "parent_child"))
            logger.info(f"Parent-Child 快取已儲存（{len(all_keys)} parents）")
        except Exception as e:
            logger.warning(f"Parent-Child 快取儲存失敗：{e}")

    def _load_or_build_parent_child(self):
        """嘗試從 pickle 載入 Parent-Child，否則全量重建"""
        pc_faiss_path = INDEX_DIR / "parent_child" / "index.faiss"
        if PARENT_CHILD_CACHE.exists() and pc_faiss_path.exists():
            try:
                with open(PARENT_CHILD_CACHE, "rb") as f:
                    data = pickle.load(f)
                self.parent_store = InMemoryStore()
                keys, values = data["keys"], data["values"]
                self.parent_store.mset(list(zip(keys, values)))
                self.pc_vectorstore = FAISS.load_local(
                    str(INDEX_DIR / "parent_child"),
                    self.embedding_model,
                    allow_dangerous_deserialization=True,
                )
                self.parent_retriever = ParentDocumentRetriever(
                    vectorstore=self.pc_vectorstore,
                    docstore=self.parent_store,
                    child_splitter=self.child_splitter,
                    parent_splitter=self.parent_splitter,
                )
                logger.info(f"Parent-Child 從快取載入（{len(keys)} parents）")
                return
            except Exception as e:
                logger.warning(f"Parent-Child 快取載入失敗，重建：{e}")

        self._build_parent_child_index()
        self._save_parent_child()

    def _build_parent_child_index(self):
        self.parent_store = InMemoryStore()
        self.pc_vectorstore = self._new_empty_faiss()
        self.parent_retriever = ParentDocumentRetriever(
            vectorstore=self.pc_vectorstore,
            docstore=self.parent_store,
            child_splitter=self.child_splitter,
            parent_splitter=self.parent_splitter,
        )

        docs = self.dm.get_documents()
        parent_docs: List[Document] = []
        for d in docs:
            full_text = self.dm.get_full_text_by_document_id(d["id"])
            if not full_text.strip():
                continue
            parent_docs.append(Document(
                page_content=full_text,
                metadata={
                    "document_id": d["id"],
                    "filename": d["original_name"],
                    "category": d["category"],
                    "tags": d.get("tags") or "",
                    "description": d.get("description") or "",
                },
            ))
        if parent_docs:
            self.parent_retriever.add_documents(parent_docs)
        logger.info(f"Parent-Child 索引建立完成，parents={len(parent_docs)}")

    def _add_parent_child_for_doc(self, document_id: int):
        if not self.parent_retriever:
            self._build_parent_child_index()
            return
        meta = self.dm.get_document_by_id(document_id)
        if not meta:
            return
        full_text = self.dm.get_full_text_by_document_id(document_id)
        if not full_text.strip():
            return
        pd_doc = Document(
            page_content=full_text,
            metadata={
                "document_id": document_id,
                "filename": meta["original_name"],
                "category": meta["category"],
                "tags": meta.get("tags") or "",
                "description": meta.get("description") or "",
            },
        )
        self.parent_retriever.add_documents([pd_doc])

    # ───── PostgreSQL tsvector FTS 候選 ─────
    def _fts_candidates(self, query: str, limit: int = 30) -> List[Dict[str, Any]]:
        """回傳 list[dict]，每項包含 content, metadata, rank"""
        if not self.dm.fts_enabled or not query.strip():
            return []
        q = self.dm._fts_safe_query(query)
        try:
            with get_cursor(SCHEMA) as cur:
                cur.execute("""
                    SELECT dc.id AS chunk_rowid,
                           dc.content,
                           d.id AS document_id,
                           dc.id AS chunk_id,
                           d.original_name,
                           d.category,
                           d.tags,
                           d.description,
                           ts_rank(to_tsvector('simple', dc.content), plainto_tsquery('simple', %s)) AS bm25_score
                    FROM document_chunks dc
                    JOIN documents d ON d.id = dc.document_id
                    WHERE d.status='active'
                      AND to_tsvector('simple', dc.content) @@ plainto_tsquery('simple', %s)
                    ORDER BY bm25_score DESC
                    LIMIT %s
                """, (q, q, int(limit)))
                rows = cur.fetchall()
            return [dict(r) for r in rows] if rows else []
        except Exception as e:
            logger.warning(f"tsvector 查詢失敗，fallback LIKE：{e}")
            like = f"%{query.lower()}%"
            with get_cursor(SCHEMA) as cur:
                cur.execute("""
                    SELECT dc.id AS chunk_rowid,
                           dc.content,
                           d.id AS document_id,
                           dc.id AS chunk_id,
                           d.original_name,
                           d.category,
                           d.tags,
                           d.description,
                           0 AS bm25_score
                    FROM document_chunks dc
                    JOIN documents d ON d.id = dc.document_id
                    WHERE d.status='active'
                      AND (LOWER(dc.content) LIKE %s OR LOWER(d.original_name) LIKE %s OR LOWER(d.tags) LIKE %s)
                    LIMIT %s
                """, (like, like, like, limit))
                rows = cur.fetchall()
            return [dict(r) for r in rows] if rows else []

    # ───── P2: Reciprocal Rank Fusion ─────
    @staticmethod
    def _rrf_merge(
        ranked_lists: List[List[Tuple[Document, float]]],
        k: int = RRF_K,
    ) -> List[Tuple[Document, float]]:
        """
        合併多個排序列表。
        每個列表為 [(Document, score), ...] 其中 score 越高越好。
        回傳按 RRF score 降序排列的 [(Document, rrf_score), ...]。
        """
        def _doc_key(d: Document) -> str:
            return f"{d.metadata.get('document_id')}:{d.metadata.get('chunk_id')}:{d.page_content[:60]}"

        scores: Dict[str, float] = {}
        doc_map: Dict[str, Document] = {}

        for ranked in ranked_lists:
            for rank, (doc, _score) in enumerate(ranked):
                key = _doc_key(doc)
                scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)
                if key not in doc_map:
                    doc_map[key] = doc

        merged = [(doc_map[key], score) for key, score in scores.items()]
        merged.sort(key=lambda x: x[1], reverse=True)
        return merged

    # ───── 壓縮檢索器 ─────
    def _build_compression_retriever(
        self,
        base_retriever,
        use_openai: bool = False,
        openai_model: str = "gpt-4o-mini",
        similarity_threshold: float = 0.5,
    ):
        emb_filter = EmbeddingsFilter(
            embeddings=self.embedding_model,
            similarity_threshold=similarity_threshold,  # P2: 提高至 0.5
        )

        base_compressor = emb_filter
        if use_openai and ChatOpenAI and LLMChainExtractor and os.environ.get("OPENAI_API_KEY"):
            try:
                llm = ChatOpenAI(model=openai_model, temperature=0)
                extractor = LLMChainExtractor.from_llm(llm)
                base_compressor = DocumentCompressorPipeline(transformers=[emb_filter, extractor])
            except Exception as e:
                logger.warning(f"LLMChainExtractor 無法建立，改用 EmbeddingsFilter：{e}")
                base_compressor = emb_filter

        return ContextualCompressionRetriever(
            base_compressor=base_compressor,
            base_retriever=base_retriever,
        )

    # ───── P5 + P0: HyDE（條件觸發 + async 修正） ─────
    async def _hyde_generate(self, question: str, use_openai: bool, openai_model: str) -> str:
        tmpl = (
            "You are a helpful assistant. Create a concise, factual, step-by-step hypothetical answer "
            "to the user's question below. Avoid placeholders. Keep it within 120~180 words.\n\n"
            f"QUESTION:\n{question}\n\nHYPOTHETICAL ANSWER:\n"
        )
        try:
            if use_openai and os.environ.get("OPENAI_API_KEY") and ChatOpenAI:
                llm = ChatOpenAI(model=openai_model, temperature=0.1, max_tokens=220)
                out = await asyncio.to_thread(llm.invoke, tmpl)
                if hasattr(out, "content"):
                    return str(out.content).strip()
                return str(out).strip()
            return await ollama_generate(tmpl, model_name=DEFAULT_OLLAMA_MODEL)
        except Exception as e:
            logger.warning(f"HyDE 生成失敗，忽略 HyDE：{e}")
            return ""

    # ───── P0+P2+P5: 查詢主流程（async + RRF + 條件 HyDE） ─────
    async def query_documents(
        self,
        question: str,
        k: int = 5,
        use_hyde: bool = True,
        use_parent: bool = True,
        use_compression: bool = True,
        use_openai: bool = False,
        openai_model: str = "gpt-4o-mini",
    ) -> List[Document]:
        if not self.vectorstore:
            return []

        ranked_lists: List[List[Tuple[Document, float]]] = []

        # ① FAISS similarity_search_with_score（帶分數）
        try:
            faiss_results = self.vectorstore.similarity_search_with_score(
                question, k=max(k * 3, 20)
            )
            # score 越小越相似（L2）或越大越相似（IP）
            # 統一為越大越好（IP 已是）
            ranked_lists.append([(doc, score) for doc, score in faiss_results])
        except Exception as e:
            logger.warning(f"FAISS 檢索失敗：{e}")
            faiss_results = []

        # ② tsvector BM25 候選
        try:
            fts_rows = self._fts_candidates(question, limit=30)
            if fts_rows:
                fts_docs = []
                for i, r in enumerate(fts_rows):
                    d = Document(
                        page_content=r["content"],
                        metadata={
                            "document_id": r["document_id"],
                            "chunk_id": r["chunk_id"],
                            "filename": r["original_name"],
                            "category": r["category"],
                            "tags": r["tags"] or "",
                            "description": r["description"] or "",
                        },
                    )
                    # BM25 score — use rank-based score
                    fts_docs.append((d, 1.0 / (i + 1)))
                ranked_lists.append(fts_docs)
        except Exception as e:
            logger.warning(f"FTS 候選合併失敗：{e}")

        # ③ P5: HyDE — 條件觸發
        # 只在文件數 >= HYDE_MIN_DOCS 且首次檢索最高分低於閾值時才啟用
        if use_hyde and faiss_results:
            best_score = faiss_results[0][1] if faiss_results else 0.0
            total_docs = len(self.dm.get_documents())
            should_hyde = total_docs >= HYDE_MIN_DOCS and best_score < HYDE_SCORE_THRESHOLD

            if should_hyde:
                hyde_txt = await self._hyde_generate(question, use_openai=use_openai, openai_model=openai_model)
                if hyde_txt:
                    try:
                        hyde_results = self.vectorstore.similarity_search_with_score(
                            hyde_txt, k=max(k, 10)
                        )
                        ranked_lists.append([(doc, score) for doc, score in hyde_results])
                    except Exception as e:
                        logger.warning(f"HyDE 檢索失敗：{e}")

        # ④ Parent-Child：擷取上位段落
        if use_parent and self.parent_retriever:
            try:
                parent_docs = self.parent_retriever.invoke(question)
                if parent_docs:
                    parent_ranked = []
                    for i, pd_doc in enumerate(parent_docs):
                        pd_doc.metadata.setdefault("document_id", pd_doc.metadata.get("document_id"))
                        pd_doc.metadata.setdefault("chunk_id", f"parent-{hash(pd_doc.page_content) & 0xffff}")
                        parent_ranked.append((pd_doc, 1.0 / (i + 1)))
                    ranked_lists.append(parent_ranked)
            except Exception as e:
                logger.warning(f"Parent-Child 檢索失敗：{e}")

        # ⑤ P2: RRF 融合所有列表
        if not ranked_lists:
            return []
        merged = self._rrf_merge(ranked_lists)

        # ⑥ Contextual Compression（對 top 結果做 post-filter）
        if use_compression and len(merged) > k:
            try:
                # 只對 merged top 結果做嵌入相似度過濾
                emb_filter = EmbeddingsFilter(
                    embeddings=self.embedding_model,
                    similarity_threshold=0.5,
                )
                top_docs = [doc for doc, _score in merged[:k * 3]]
                filtered = emb_filter.compress_documents(top_docs, question)
                if filtered:
                    merged = [(d, 1.0) for d in filtered]
            except Exception as e:
                logger.warning(f"Compression 過濾失敗，使用 RRF 結果：{e}")

        # ⑦ 取 top-k
        final_docs = [doc for doc, _score in merged[:max(k, 5)]]
        return final_docs


# ═══════════════════════════════════════════════════════════════════
# 生成模型（Ollama / OpenAI）— P0: 正確 async
# ═══════════════════════════════════════════════════════════════════
def _strip_thinking(text: str) -> str:
    cleaned = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    lines = [ln for ln in cleaned.splitlines() if not ln.strip().startswith(("Thinking:", "Processing:", "Analyzing:"))]
    return "\n".join(lines).strip()


async def ollama_generate(prompt: str, model_name: str = DEFAULT_OLLAMA_MODEL, stream: bool = False) -> str:
    """本地 Ollama 生成；num_ctx 設 8192、keep_alive 15m"""
    try:
        payload = {
            "model": model_name,
            "prompt": prompt,
            "stream": stream,
            "options": {
                "temperature": 0.2,
                "top_p": 0.1,
                "top_k": 10,
                "repeat_penalty": 1.1,
                "num_predict": 2000,
                "num_ctx": 8192,
            },
            "keep_alive": "15m",
        }
        r = await asyncio.to_thread(requests.post, "http://localhost:11434/api/generate", json=payload, timeout=120)
        r.raise_for_status()
        result = r.json()
        return _strip_thinking(result.get("response", ""))
    except requests.exceptions.RequestException as e:
        raise Exception(f"Ollama 連線錯誤: {str(e)}")
    except Exception as e:
        raise Exception(f"Ollama 回應處理錯誤: {str(e)}")


async def openai_generate(prompt: str, model_name: str = "gpt-4o-mini") -> str:
    """OpenAI Chat Completions"""
    try:
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise Exception("OpenAI API key not found. Please set OPENAI_API_KEY environment variable.")

        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": model_name,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant specializing in document analysis, SOP guidance, and "
                        "process improvement. Provide clear, actionable responses based on the provided documents. "
                        "If the provided documents do not contain relevant information, clearly state that."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            "max_tokens": 2000,
            "temperature": 0.2,
        }
        r = await asyncio.to_thread(
            requests.post, "https://api.openai.com/v1/chat/completions",
            headers=headers, json=payload, timeout=60,
        )
        if r.status_code == 200:
            data = r.json()
            return data["choices"][0]["message"]["content"].strip()
        else:
            detail = r.json().get("error", {}).get("message", r.text)
            raise Exception(f"OpenAI API error: {r.status_code} - {detail}")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Error connecting to OpenAI API: {str(e)}")
    except Exception as e:
        raise Exception(f"Error processing OpenAI request: {str(e)}")


# ═══════════════════════════════════════════════════════════════════
# P4: Prompt 模板（強化 anti-hallucination + 結構化輸出）
# ═══════════════════════════════════════════════════════════════════
PROMPT_TEMPLATES = {
    "sop_query": PromptTemplate(
        input_variables=["context", "question"],
        template="""You are an SOP assistant for a manufacturing company. Answer ONLY based on the provided documents.

IMPORTANT RULES:
- If the documents do NOT contain relevant information, say "文件中未找到相關資訊" and suggest what to search for.
- Do NOT make up procedures or steps that are not in the documents.
- Cite the source document name for each answer.

DOCUMENTS:
{context}

QUESTION: {question}

Please answer in the following format:
## 回答
[Your answer based on the documents]

## 步驟（if applicable）
1. ...
2. ...

## 注意事項（if applicable）
- ...

## 來源文件
- [Document names used]"""
    ),
    "form_assistance": PromptTemplate(
        input_variables=["context", "question"],
        template="""You are a form assistance expert. Answer ONLY based on the provided documents.

IMPORTANT: If the documents do NOT contain the form or relevant information, say "文件中未找到相關表格" and do NOT fabricate form fields.

DOCUMENTS:
{context}

QUESTION: {question}

Please provide:
## 表格指導
[Clear guidance on how to fill out or use the forms]

## 必填欄位（if identifiable）
- ...

## 來源文件
- [Document names used]"""
    ),
    "improvement_analysis": PromptTemplate(
        input_variables=["context", "question"],
        template="""Analyze the following documents for improvement opportunities.

IMPORTANT: Base your analysis ONLY on the provided documents. If the documents lack sufficient detail for analysis, state this clearly.

DOCUMENTS:
{context}

QUESTION: {question}

Provide your analysis in this format:
## 現狀分析
[Current state based on documents]

## 發現的問題
1. ...

## 改進建議
1. [Suggestion] — 可行性: [高/中/低]
2. ...

## 來源文件
- [Document names used]"""
    ),
    "document_request": PromptTemplate(
        input_variables=["context", "question", "found_documents"],
        template="""The user is requesting a specific document.

FOUND DOCUMENTS:
{found_documents}

RELATED CONTEXT:
{context}

USER REQUEST: {question}

Identify which document(s) match the request and provide download guidance. If no match is found, suggest alternative search terms."""
    ),
    "general_document": PromptTemplate(
        input_variables=["context", "question"],
        template="""You are a document analysis assistant. Answer ONLY based on the provided documents.

CRITICAL RULES:
- If the documents do NOT contain relevant information to answer the question, respond with: "文件中未找到相關資訊，建議上傳相關文件或嘗試不同的關鍵詞。"
- Do NOT generate information that is not present in the documents.
- Always cite the source document name.

DOCUMENTS:
{context}

QUESTION: {question}

## 回答
[Your answer]

## 來源文件
- [Document names used]"""
    ),
}

# ─────────────────────────── 查詢類型判斷 / 名稱解析 ───────────────────────────
def determine_query_type(question: str) -> tuple[str, bool]:
    q = question.lower()
    doc_req_kw = ['我需要', 'i need', '給我', 'give me', '下載', 'download',
                  '導出', 'export', '找到', 'find', '提供', 'provide',
                  '發送', 'send', '取得', 'get', '要', 'want']
    doc_type_kw = ['文件', '文檔', '表格', '表單', 'form', 'document',
                   'file', 'template', '模板', 'sop', '程序']
    is_doc_req = any(k in q for k in doc_req_kw) and any(k in q for k in doc_type_kw)
    if is_doc_req:
        return "document_request", True

    sop_kw = ['sop', 'procedure', 'process', 'step', 'how to', '流程', '程序', '步驟']
    form_kw = ['form', 'template', 'fill', 'complete', '表格', '填寫', '模板']
    improve_kw = ['improve', 'optimize', 'better', 'enhance', 'suggestion', '改進', '優化', '建議']

    if any(k in q for k in improve_kw):
        return "improvement_analysis", False
    elif any(k in q for k in sop_kw):
        return "sop_query", False
    elif any(k in q for k in form_kw):
        return "form_assistance", False
    else:
        return "general_document", False


def extract_document_name_from_question(question: str) -> str:
    q = question.strip()
    m = re.search(r'["""](.+?)["""]', q)
    if m:
        return m.group(1).strip()
    m = re.search(r'[（(](.+?)[)）]', q)
    if m:
        return m.group(1).strip()

    remove_words = ['我需要', '給我', '下載', '導出', '找到', '提供', '發送', '取得', '要',
                    'i need', 'give me', 'download', 'export', 'find', 'provide', 'send', 'get', 'want',
                    '這個', '那個', '的', 'this', 'that', 'the', 'a', 'an']
    lower_q = q.lower()
    for w in remove_words:
        lower_q = lower_q.replace(w, ' ')
    keywords = ['文件', '文檔', '表格', '表單', 'form', 'document', 'file', 'template', '模板',
                'sop', '程序', 'checklist', '清單', 'manual', '手冊', 'policy', '政策']
    for kw in keywords:
        if kw in lower_q:
            parts = lower_q.split(kw, 1)
            if len(parts) > 1 and parts[1].strip():
                return parts[1].strip()
    return " ".join(lower_q.split())


# ═══════════════════════════════════════════════════════════════════
# Orchestrator：AI 文檔分析 — P0: async 修正
# ═══════════════════════════════════════════════════════════════════
class AIDocumentAnalytics:
    def __init__(self):
        self.document_manager = DocumentManager()
        self.rag_system = DocumentRAG(self.document_manager)
        self.prompt_templates = PROMPT_TEMPLATES
        logger.info("AI Document Analytics initialized")

    def upload_document(
        self, file_path: str, original_name: str, category: str,
        uploaded_by: str, tags: str = "", description: str = ""
    ) -> Dict[str, Any]:
        result = self.document_manager.upload_document(
            file_path, original_name, category, uploaded_by, tags, description
        )
        if result.get("success"):
            self.rag_system.add_document_chunks(result["document_id"])
            logger.info(f"Document uploaded & indexes updated: {original_name}")
        return result

    async def query_documents(
        self, question: str, use_openai: bool = False, openai_model: str = "gpt-4o-mini"
    ) -> Dict[str, Any]:
        try:
            query_type, is_doc_req = determine_query_type(question)

            # 文件請求優先以名稱搜尋
            if is_doc_req:
                search_term = extract_document_name_from_question(question)
                found = self.document_manager.search_documents_by_name(search_term)
                if found:
                    first = found[0]
                    resp = "找到了您需要的文檔：\n\n"
                    resp += f"**文檔名稱**: {first['original_name']}\n"
                    resp += f"**類別**: {DOCUMENT_CATEGORIES.get(first['category'], first['category'])}\n"
                    resp += f"**文件類型**: {first['file_type'].upper()}\n"
                    if first.get('description'):
                        resp += f"**描述**: {first['description']}\n"
                    if first.get('tags'):
                        resp += f"**標籤**: {first['tags']}\n"
                    resp += f"\n📥 **下載文檔**: [下載 {first['original_name']}](/api/ai/documents/{first['id']}/download)\n"

                    if len(found) > 1:
                        resp += f"\n📚 還找到其他 {len(found) - 1} 個相關文檔：\n"
                        for doc in found[1:4]:
                            resp += f"- [{doc['original_name']}](/api/ai/documents/{doc['id']}/download) ({doc['file_type'].upper()})\n"

                    return {
                        "answer": resp,
                        "source_documents": [{
                            "filename": d['original_name'],
                            "category": d['category'],
                            "document_id": d['id'],
                            "content_preview": f"Document type: {d['file_type'].upper()}",
                        } for d in found[:5]],
                        "status": "success",
                        "ai_provider": "system",
                        "query_type": "document_request",
                        "documents_found": found,
                    }
                else:
                    query_type = "general_document"

            # RAG：FAISS + FTS + RRF + 條件 HyDE + Parent-Child + 壓縮
            relevant = await self.rag_system.query_documents(
                question=question,
                k=8,
                use_hyde=True,
                use_parent=True,
                use_compression=True,
                use_openai=use_openai,
                openai_model=openai_model,
            )
            if not relevant:
                return {
                    "answer": "抱歉，找不到相關文檔內容可用於回答。請確認已上傳相關文件，或更換關鍵詞再試。",
                    "source_documents": [],
                    "status": "no_documents",
                }

            # P4: 構建帶來源標記的上下文
            context_parts = []
            for i, d in enumerate(relevant):
                source = d.metadata.get('filename', 'Unknown')
                category = d.metadata.get('category', '')
                chunk_idx = d.metadata.get('chunk_index', '?')
                context_parts.append(
                    f"[Source #{i + 1}: {source}, Category: {category}, Section: {chunk_idx}]\n{d.page_content}"
                )
            context = "\n\n---\n\n".join(context_parts)

            tmpl = self.prompt_templates.get(query_type, self.prompt_templates["general_document"])
            if query_type == "document_request":
                prompt = tmpl.format(context=context, question=question, found_documents="")
            else:
                prompt = tmpl.format(context=context, question=question)

            # P0: 正確 await async 生成函式
            if use_openai:
                answer = await openai_generate(prompt, openai_model)
                provider = "openai"
            else:
                answer = await ollama_generate(prompt, model_name=DEFAULT_OLLAMA_MODEL)
                provider = "ollama"

            src = [{
                "filename": d.metadata.get("filename"),
                "category": d.metadata.get("category"),
                "document_id": d.metadata.get("document_id"),
                "content_preview": (d.page_content[:200] + "...") if len(d.page_content) > 200 else d.page_content,
            } for d in relevant]

            return {
                "answer": answer,
                "source_documents": src,
                "status": "success",
                "ai_provider": provider,
                "query_type": query_type,
            }
        except Exception as e:
            logger.exception("Query processing error")
            return {"answer": f"處理查詢時出錯: {str(e)}", "source_documents": [], "status": "error"}

    def get_documents(self, category: str = None) -> List[Dict[str, Any]]:
        return self.document_manager.get_documents(category)

    def get_document_by_id(self, document_id: int) -> Optional[Dict[str, Any]]:
        return self.document_manager.get_document_by_id(document_id)

    def get_document_path(self, document_id: int) -> Optional[Path]:
        return self.document_manager.get_document_path(document_id)

    def delete_document(self, document_id: int) -> bool:
        success = self.document_manager.delete_document(document_id)
        if success:
            self.rag_system.remove_document(document_id)
            logger.info(f"Document deleted & indexes rebuilt(excluded): {document_id}")
        return success

    def get_categories(self) -> Dict[str, str]:
        return DOCUMENT_CATEGORIES

    def get_system_status(self) -> Dict[str, Any]:
        try:
            docs = self.document_manager.get_documents()
            with get_cursor(SCHEMA) as cur:
                cur.execute("SELECT COUNT(*) AS c FROM document_chunks")
                chunk_count = cur.fetchone()["c"]
            faiss_path = INDEX_DIR / "index.faiss"
            index_ready = self.rag_system.vectorstore is not None
            index_size = faiss_path.stat().st_size if faiss_path.exists() else 0

            cat = {}
            for d in docs:
                cat[d["category"]] = cat.get(d["category"], 0) + 1

            return {
                "status": "ready",
                "total_documents": len(docs),
                "total_chunks": int(chunk_count),
                "categories": cat,
                "vector_store_ready": index_ready,
                "vector_store_size_bytes": index_size,
                "supported_file_types": sorted(list(SUPPORTED_EXTENSIONS)),
                "ollama_model": DEFAULT_OLLAMA_MODEL,
                "embedding_model": EMBEDDING_MODEL,
                "openai_available": bool(os.environ.get("OPENAI_API_KEY")),
                "fts_enabled": self.document_manager.fts_enabled,
                "parent_child_enabled": self.rag_system.parent_retriever is not None,
                "hyde_threshold": HYDE_SCORE_THRESHOLD,
                "hyde_min_docs": HYDE_MIN_DOCS,
                "rrf_k": RRF_K,
            }
        except Exception as e:
            logger.exception("Error getting system status")
            return {"status": "error", "message": str(e)}


# ─────────────────────────── 全局 AI 引擎實例 ───────────────────────────
_ai_engine_instance: Optional[AIDocumentAnalytics] = None
_ai_engine_lock = threading.Lock()


def get_ai_engine() -> AIDocumentAnalytics:
    """Get singleton AI engine and initialize it lazily on first use."""
    global _ai_engine_instance
    if _ai_engine_instance is None:
        with _ai_engine_lock:
            if _ai_engine_instance is None:
                logger.info("Initializing AI Document Analytics engine on first use")
                _ai_engine_instance = AIDocumentAnalytics()
    return _ai_engine_instance


class _LazyAIEngine:
    """Proxy object that defers heavy AI engine initialization."""

    def __getattr__(self, attr):
        return getattr(get_ai_engine(), attr)


ai_engine = _LazyAIEngine()
